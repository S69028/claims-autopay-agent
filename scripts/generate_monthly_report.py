#!/usr/bin/env python3
"""Generate a monthly auto-insurance operations report as a single DOCX document.

The report is built from either Supabase REST or local simulation CSVs and is
intended to be a single document suitable for archive/download workflows.
"""

from __future__ import annotations

import argparse
import calendar
import csv
import json
import math
import os
import sys
import urllib.error
import urllib.parse
import urllib.request
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable


ROOT = Path(__file__).resolve().parent.parent
DEPS = ROOT / ".deps"
if DEPS.exists():
    sys.path.insert(0, str(DEPS))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SUPABASE_URL = os.environ.get("SUPABASE_URL", "https://ejgruldmrfbtalplwowr.supabase.co")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "sb_publishable_1swNgizDHWX6Dh7bwJ2WTg_t1qKUvBD")
DEFAULT_REPORT_MONTH = "2025-12"
DEFAULT_OUT = ROOT / "reports" / "monthly_auto_payment_report_2025-12.docx"
FONT_BODY = "바탕체"
FONT_SYMBOL = "Segoe UI Symbol"
BLACK = "000000"
RED = "ff0000"
BLUE = "2400FF"


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def month_prev(month: str) -> str:
    year, mm = month.split("-")
    y = int(year)
    m = int(mm)
    if m == 1:
        return f"{y - 1}-12"
    return f"{y}-{m - 1:02d}"


def as_int(value: str | int | float | None) -> int:
    if value in (None, "", "-"):
        return 0
    if isinstance(value, int):
        return value
    if isinstance(value, float):
        return int(round(value))
    return int(float(str(value).replace(",", "")))


def as_float(value: str | int | float | None) -> float:
    if value in (None, "", "-"):
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)
    return float(str(value).replace(",", ""))


def pct(numerator: int, denominator: int) -> float | None:
    if denominator == 0:
        return None
    return round((numerator / denominator) * 100, 2)


def pct_change(current: float, previous: float) -> float | None:
    if previous == 0:
        return None
    return round(((current - previous) / previous) * 100, 2)


def fmt_int(value: int) -> str:
    return f"{value:,}"


def fmt_pct(value: float | None, digits: int = 1, suffix: str = "%") -> str:
    if value is None:
        return "미계산"
    return f"{value:.{digits}f}{suffix}"


def fmt_pp(value: float | None, digits: int = 1) -> str:
    if value is None:
        return "미계산"
    sign = "+" if value >= 0 else ""
    return f"{sign}{value:.{digits}f}pp"


def fmt_change(count_delta: int, pct_delta: float | None) -> str:
    if pct_delta is None:
        if count_delta == 0:
            return "0 (0.0%)"
        return f"{fmt_int(count_delta)} (신규)"
    sign = "+" if pct_delta >= 0 else ""
    return f"{fmt_int(count_delta)} ({sign}{pct_delta:.1f}%)"


def http_get_json(path: str) -> list[dict]:
    url = f"{SUPABASE_URL}{path}"
    request = urllib.request.Request(
        url,
        headers={
            "apikey": SUPABASE_KEY,
            "Authorization": f"Bearer {SUPABASE_KEY}",
            "Accept": "application/json",
        },
    )
    with urllib.request.urlopen(request, timeout=30) as response:
        text = response.read().decode("utf-8")
    return json.loads(text)


def load_from_supabase(report_month: str) -> tuple[list[dict], list[dict], list[dict], list[dict], list[dict]]:
    snapshots = http_get_json("/rest/v1/monthly_auto_payment_snapshot_fact?select=*&order=snapshot_month.asc")
    current_claims = http_get_json(
        f"/rest/v1/claim_payment_fact?select=*&claim_month=eq.{report_month}&order=claim_id.asc"
    )
    previous_claims = http_get_json(
        f"/rest/v1/claim_payment_fact?select=*&claim_month=eq.{month_prev(report_month)}&order=claim_id.asc"
    )
    auto_types = http_get_json("/rest/v1/auto_payment_type_definition?select=type_code,type_name,priority_order&order=priority_order.asc")
    exclusion_types = http_get_json("/rest/v1/auto_payment_exclusion_type_definition?select=type_code,type_name,priority_order&order=priority_order.asc")
    return snapshots, current_claims, previous_claims, auto_types, exclusion_types


def load_from_local(report_month: str) -> tuple[list[dict], list[dict], list[dict], list[dict], list[dict]]:
    snapshots = read_csv_rows(ROOT / "data" / "simulation" / "monthly_auto_payment_snapshot_fact_2025.csv")
    current_claims = read_csv_rows(ROOT / "data" / "simulation" / f"claim_payment_fact_{report_month}.csv")
    previous_claims = read_csv_rows(ROOT / "data" / "simulation" / f"claim_payment_fact_{month_prev(report_month)}.csv")
    auto_types = read_csv_rows(ROOT / "data" / "poc" / "03_auto_payment_type_definition_sample.csv")
    exclusion_types = read_csv_rows(ROOT / "data" / "poc" / "06_auto_payment_exclusion_type_definition_sample.csv")
    return snapshots, current_claims, previous_claims, auto_types, exclusion_types


def load_data(source: str, report_month: str) -> tuple[list[dict], list[dict], list[dict], list[dict], list[dict]]:
    if source == "local":
        return load_from_local(report_month)
    if source == "db":
        return load_from_supabase(report_month)
    try:
        return load_from_supabase(report_month)
    except Exception:
        return load_from_local(report_month)


def confirmed_snapshot(snapshots: list[dict], report_month: str) -> tuple[dict, dict]:
    ordered = sorted(snapshots, key=lambda row: row["snapshot_month"])
    latest = next((row for row in reversed(ordered) if row.get("data_frozen_flag") == "Y"), ordered[-1])
    idx = ordered.index(latest)
    previous = ordered[idx - 1] if idx > 0 else latest
    if latest["snapshot_month"] != report_month:
        match = next((row for row in ordered if row["snapshot_month"] == report_month), None)
        if match is not None:
            latest = match
            idx = ordered.index(match)
            previous = ordered[idx - 1] if idx > 0 else match
    return latest, previous


def group_count(rows: list[dict], key: str, predicate=None) -> Counter:
    counter: Counter[str] = Counter()
    for row in rows:
        if predicate is not None and not predicate(row):
            continue
        counter[str(row.get(key, "")).strip() or "-"] += 1
    return counter


def map_type_names(def_rows: list[dict]) -> dict[str, str]:
    return {str(row.get("type_code", "")).strip(): str(row.get("type_name", "")).strip() for row in def_rows}


def top_items(counter: Counter, limit: int = 5) -> list[tuple[str, int]]:
    return counter.most_common(limit)


def build_context(
    snapshots: list[dict],
    current_claims: list[dict],
    previous_claims: list[dict],
    auto_types: list[dict],
    exclusion_types: list[dict],
    report_month: str,
) -> dict:
    latest, previous = confirmed_snapshot(snapshots, report_month)
    prev_month = previous["snapshot_month"] if previous else month_prev(report_month)
    auto_type_names = map_type_names(auto_types)
    exclusion_type_names = map_type_names(exclusion_types)

    current_auto = [row for row in current_claims if row.get("auto_payment_decision") == "AUTO_PAY"]
    current_excluded = [row for row in current_claims if row.get("auto_payment_decision") == "EXCLUDE"]
    current_manual = [row for row in current_claims if row.get("auto_payment_decision") == "MANUAL_REVIEW"]
    previous_auto = [row for row in previous_claims if row.get("auto_payment_decision") == "AUTO_PAY"]
    previous_excluded = [row for row in previous_claims if row.get("auto_payment_decision") == "EXCLUDE"]
    previous_manual = [row for row in previous_claims if row.get("auto_payment_decision") == "MANUAL_REVIEW"]

    channel_total = group_count(current_claims, "receipt_channel")
    channel_auto = group_count(current_auto, "receipt_channel")
    prev_channel_total = group_count(previous_claims, "receipt_channel")
    prev_channel_auto = group_count(previous_auto, "receipt_channel")
    treatment_auto = group_count(current_auto, "treatment_type")
    prev_treatment_auto = group_count(previous_auto, "treatment_type")
    top_auto_types = top_items(group_count(current_auto, "auto_payment_type_code"))
    top_exclusion_types = top_items(group_count(current_excluded, "auto_payment_exclusion_code"))

    prev_auto = group_count(previous_claims, "auto_payment_type_code", lambda row: row.get("auto_payment_decision") == "AUTO_PAY")
    prev_exclusion = group_count(previous_claims, "auto_payment_exclusion_code", lambda row: row.get("auto_payment_decision") == "EXCLUDE")
    prev_auto_map = {k: v for k, v in prev_auto.items()}
    prev_exclusion_map = {k: v for k, v in prev_exclusion.items()}

    segment_pairs = Counter(
        (
            str(row.get("segment_before", "")).strip() or "-",
            str(row.get("segment_after", "")).strip() or "-",
            str(row.get("segment_change_reason", "")).strip() or "원인 확인 필요",
        )
        for row in current_claims
        if str(row.get("segment_change_flag", "")).upper() == "Y"
    )

    channel_rows = []
    for channel, total in sorted(channel_total.items(), key=lambda item: (-item[1], item[0])):
        auto = channel_auto.get(channel, 0)
        share = pct(auto, total)
        prev_total = prev_channel_total.get(channel, 0)
        prev_auto = prev_channel_auto.get(channel, 0)
        prev_share = pct(prev_auto, prev_total)
        channel_rows.append(
            {
                "channel": channel,
                "total": total,
                "prev_total": prev_total,
                "auto": auto,
                "prev_auto": prev_auto,
                "share": share,
                "prev_share": prev_share,
                "delta_share": None if share is None or prev_share is None else round(share - prev_share, 2),
            }
        )

    treatment_rows = []
    total_auto = len(current_auto)
    for treatment, count in sorted(treatment_auto.items(), key=lambda item: (-item[1], item[0])):
        prev_count = prev_treatment_auto.get(treatment, 0)
        prev_total_auto = len(previous_auto)
        current_share = pct(count, total_auto)
        previous_share = pct(prev_count, prev_total_auto)
        treatment_rows.append(
            {
                "treatment": treatment,
                "count": count,
                "prev_count": prev_count,
                "share": current_share,
                "prev_share": previous_share,
                "delta_share": None if current_share is None or previous_share is None else round(current_share - previous_share, 2),
            }
        )

    auto_type_rows = []
    for idx, (code, count) in enumerate(top_auto_types, 1):
        prev_count = prev_auto_map.get(code, 0)
        auto_type_rows.append(
            {
                "rank": idx,
                "code": code,
                "name": auto_type_names.get(code, code),
                "count": count,
                "share": pct(count, len(current_auto)),
                "prev_count": prev_count,
                "delta_count": count - prev_count,
                "delta_pct": pct_change(count, prev_count),
            }
        )

    exclusion_rows = []
    for idx, (code, count) in enumerate(top_exclusion_types, 1):
        prev_count = prev_exclusion_map.get(code, 0)
        exclusion_rows.append(
            {
                "rank": idx,
                "code": code,
                "name": exclusion_type_names.get(code, code),
                "count": count,
                "share": pct(count, len(current_excluded)),
                "prev_count": prev_count,
                "delta_count": count - prev_count,
                "delta_pct": pct_change(count, prev_count),
            }
        )

    segment_rows = []
    for (before, after, reason), count in segment_pairs.most_common(5):
        segment_rows.append(
            {
                "before": before,
                "after": after,
                "count": count,
                "share": pct(count, len(current_claims)),
                "reason": reason,
            }
        )

    summary_lines = [
        f"{report_month} 기준 자동지급률은 {fmt_pct(as_float(latest['auto_payment_rate']), 2)}이고 처리효율은 {fmt_pct(as_float(latest['processing_efficiency']), 2)}입니다.",
        f"전월 대비 자동지급률은 {fmt_pp(as_float(latest['auto_payment_rate_change']), 2)}, 처리효율은 {fmt_pp(as_float(latest['processing_efficiency_change']), 2)}로 집계되었습니다.",
    ]
    if str(latest.get("status_label")) == "안정":
        summary_lines.append("segment 변화는 일부 있지만 전체 상태는 허용 범위 내로 판단되어 안정 상태로 표시했습니다.")
    else:
        summary_lines.append("이번 월은 segment 또는 비율 변화가 커서 운영 해석을 함께 확인할 필요가 있습니다.")

    next_actions = [
        "전월 대비 자동지급률 상승을 주도한 유형의 분포를 재확인한다.",
        "segment 변경 사유가 `원인 확인 필요`로 남은 건은 운영관리자가 추가 확인한다.",
    ]
    if current_manual:
        next_actions.append("인심사 건수와 제외 건수가 늘어난 채널은 별도 모니터링 대상으로 둔다.")

    top_channel = max(channel_rows, key=lambda item: item["share"] or 0) if channel_rows else None
    top_treatment = max(treatment_rows, key=lambda item: item["share"] or 0) if treatment_rows else None
    top_auto = auto_type_rows[0] if auto_type_rows else None
    top_exclusion = exclusion_rows[0] if exclusion_rows else None

    overall_note = summary_lines[0]
    status_note = summary_lines[2]
    channel_note = (
        f"접수채널 중 {top_channel['channel']} 채널의 자동지급 비율이 가장 높고, "
        f"해당 채널의 전환 흐름이 전체 처리현황을 주도합니다."
        if top_channel
        else "접수채널별 비율을 확인할 데이터가 부족합니다."
    )
    treatment_note = (
        f"진료구분 중 {top_treatment['treatment']}이 자동지급 건의 중심이며 비중은 {fmt_pct(top_treatment['share'], 1) if top_treatment['share'] is not None else '미계산'}입니다."
        if top_treatment
        else "진료구분별 비율을 확인할 데이터가 부족합니다."
    )
    type_note = (
        f"상위 자동지급 유형은 {top_auto['name']}({top_auto['code']})이며 전월 대비 {fmt_change(top_auto['delta_count'], top_auto['delta_pct'])}입니다."
        if top_auto
        else "자동지급 유형을 확인할 데이터가 부족합니다."
    )
    exclusion_note = (
        f"상위 제외 유형은 {top_exclusion['name']}({top_exclusion['code']})이며 전월 대비 {fmt_change(top_exclusion['delta_count'], top_exclusion['delta_pct'])}입니다."
        if top_exclusion
        else "자동지급 제외 유형을 확인할 데이터가 부족합니다."
    )

    return {
        "latest": latest,
        "previous": previous,
        "prev_month": prev_month,
        "channel_rows": channel_rows,
        "treatment_rows": treatment_rows,
        "auto_type_rows": auto_type_rows,
        "exclusion_rows": exclusion_rows,
        "segment_rows": segment_rows,
        "summary_lines": summary_lines,
        "next_actions": next_actions,
        "current_claims": current_claims,
        "current_auto_count": len(current_auto),
        "current_excluded_count": len(current_excluded),
        "current_manual_count": len(current_manual),
        "total_claims": len(current_claims),
        "auto_type_names": auto_type_names,
        "exclusion_type_names": exclusion_type_names,
        "overall_note": overall_note,
        "status_note": status_note,
        "channel_note": channel_note,
        "treatment_note": treatment_note,
        "type_note": type_note,
        "exclusion_note": exclusion_note,
        "previous_auto_count": len(previous_auto),
        "previous_excluded_count": len(previous_excluded),
        "previous_manual_count": len(previous_manual),
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_doc_defaults(document: Document) -> None:
    styles = document.styles
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        if style_name in styles:
            style = styles[style_name]
            try:
                style.font.name = FONT_BODY
                style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
            except Exception:
                pass
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def apply_run_style(run, *, size: int, bold: bool = False, font_name: str = FONT_BODY, color: str = BLACK) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def value_color(text: str) -> str:
    return BLACK


def apply_paragraph_style(
    paragraph,
    *,
    align=None,
    line_spacing=None,
    first_line_indent=None,
    left_indent=None,
) -> None:
    if align is not None:
        paragraph.alignment = align
    fmt = paragraph.paragraph_format
    if line_spacing is not None:
        fmt.line_spacing = line_spacing
    if first_line_indent is not None:
        fmt.first_line_indent = first_line_indent
    if left_indent is not None:
        fmt.left_indent = left_indent


def add_text_paragraph(
    document: Document,
    parts: list[tuple[str, dict]],
    *,
    align=None,
    size: int = 12,
    bold: bool = False,
    line_spacing=None,
    first_line_indent=None,
    left_indent=None,
) -> None:
    p = document.add_paragraph()
    apply_paragraph_style(
        p,
        align=align,
        line_spacing=line_spacing,
        first_line_indent=first_line_indent,
        left_indent=left_indent,
    )
    for text, opts in parts:
        run = p.add_run(text)
        apply_run_style(
            run,
            size=opts.get("size", size),
            bold=opts.get("bold", bold),
            font_name=opts.get("font_name", FONT_BODY),
            color=opts.get("color", BLACK),
        )
        if opts.get("underline"):
            run.font.underline = True


def add_blank_paragraph(document: Document) -> None:
    document.add_paragraph()


def add_table(
    document: Document,
    headers: list[str],
    rows: Iterable[Iterable[object]],
    *,
    header_size: int = 12,
    row_size: int = 12,
    header_bold: bool = True,
    row_aligns: list[WD_ALIGN_PARAGRAPH] | None = None,
    row_left_align_col: int | None = None,
) -> None:
    rows = list(rows)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = str(header)
        for paragraph in hdr[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                apply_run_style(run, size=header_size, bold=header_bold)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    color = value_color(run.text)
                    apply_run_style(run, size=row_size, bold=False, color=color)
        if row_left_align_col is not None and 0 <= row_left_align_col < len(cells):
            for paragraph in cells[row_left_align_col].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return None


def add_section_heading(document: Document, title: str, subtitle: str | None = None, level: int = 1) -> None:
    p = document.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.name = FONT_BODY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(BLACK)
    if subtitle:
        p = document.add_paragraph()
        run = p.add_run(subtitle)
        apply_run_style(run, size=12)


def add_colored_paragraph(document: Document, parts: list[tuple[str, str | None]], *, align=WD_ALIGN_PARAGRAPH.LEFT, bold: bool = False, size: int = 10) -> None:
    mapped = [(text, {"color": BLACK, "size": size, "bold": bold, "font_name": FONT_BODY}) for text, _ in parts]
    add_text_paragraph(document, mapped, align=align, size=size, bold=bold)


def add_note_paragraph(document: Document, text: str) -> None:
    add_text_paragraph(
        document,
        [
            ("- ", {}),
            (text, {}),
        ],
        size=12,
        line_spacing=1.5,
        first_line_indent=Inches(0.167),
    )


def add_type_table(document: Document, rows: list[dict]) -> None:
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["순위", "코드", "유형명", "비중", "당월", "전월", "증감"]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
        for paragraph in table.cell(0, idx).paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                apply_run_style(run, size=10, bold=True)
    for row in rows:
        cells = table.add_row().cells
        values = [
            row["rank"],
            row["code"],
            row["name"],
            fmt_pct(row["share"], 1) if row["share"] is not None else "미계산",
            fmt_int(row["count"]),
            fmt_int(row["prev_count"]),
            fmt_change(row["delta_count"], row["delta_pct"]),
        ]
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    apply_run_style(run, size=10, bold=False, color=value_color(run.text))
        for paragraph in cells[2].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def build_document(ctx: dict) -> Document:
    doc = Document()
    set_doc_defaults(doc)
    year, month = ctx["latest"]["snapshot_month"].split("-")
    yy = year[-2:]
    month_label = f"'{yy}.{int(month):02d}月"
    period_start = f"'{yy}.{int(month):02d}.1"
    last_day = calendar.monthrange(int(year), int(month))[1]
    period_end = f"'{yy}.{int(month):02d}.{last_day:02d}"

    add_text_paragraph(
        doc,
        [("자동지급 현황 월간 보고", {"bold": True, "underline": True})],
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=24,
        bold=True,
    )
    add_text_paragraph(
        doc,
        [(month_label, {})],
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        size=12,
    )
    add_text_paragraph(
        doc,
        [("보험금심사팀", {})],
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        size=12,
    )

    add_text_paragraph(
        doc,
        [
            ("1. ", {"bold": True}),
            ("요 약", {"bold": True}),
        ],
        size=12,
        bold=True,
    )
    add_text_paragraph(
        doc,
        [
            ("☐", {"font_name": FONT_SYMBOL}),
            (" ", {}),
            ("전체 처리현황", {}),
        ],
        size=12,
    )
    add_text_paragraph(
        doc,
        [("- 기 간:", {}), (" ", {}), (period_start, {}), (" ~", {}), (" ", {}), (period_end, {}), (" (1개월)", {})],
        size=12,
        first_line_indent=Inches(0.167),
    )
    add_note_paragraph(doc, ctx["overall_note"])
    add_note_paragraph(doc, ctx["status_note"])
    add_table(
        doc,
        ["항목", "당월", "전월比"],
        [
            [
                "총 청구건수",
                fmt_int(as_int(ctx["latest"]["total_claim_count"])),
                fmt_change(as_int(ctx["latest"]["total_claim_count"]) - as_int(ctx["previous"]["total_claim_count"]), pct_change(as_int(ctx["latest"]["total_claim_count"]), as_int(ctx["previous"]["total_claim_count"]))),
            ],
            [
                "자동지급 대상건수",
                fmt_int(as_int(ctx["latest"]["auto_payment_candidate_count"])),
                fmt_change(as_int(ctx["latest"]["auto_payment_candidate_count"]) - as_int(ctx["previous"]["auto_payment_candidate_count"]), pct_change(as_int(ctx["latest"]["auto_payment_candidate_count"]), as_int(ctx["previous"]["auto_payment_candidate_count"]))),
            ],
            [
                "자동지급 건수",
                fmt_int(as_int(ctx["latest"]["auto_payment_count"])),
                fmt_change(as_int(ctx["latest"]["auto_payment_count"]) - as_int(ctx["previous"]["auto_payment_count"]), pct_change(as_int(ctx["latest"]["auto_payment_count"]), as_int(ctx["previous"]["auto_payment_count"]))),
            ],
            [
                "제외/인심사 건수",
                fmt_int(as_int(ctx["latest"]["exclusion_count"]) + as_int(ctx["latest"]["manual_review_count"])),
                fmt_change(
                    (as_int(ctx["latest"]["exclusion_count"]) + as_int(ctx["latest"]["manual_review_count"]))
                    - (as_int(ctx["previous"]["exclusion_count"]) + as_int(ctx["previous"]["manual_review_count"])),
                    pct_change(
                        as_int(ctx["latest"]["exclusion_count"]) + as_int(ctx["latest"]["manual_review_count"]),
                        as_int(ctx["previous"]["exclusion_count"]) + as_int(ctx["previous"]["manual_review_count"]),
                    ),
                ),
            ],
            ["자동지급률", fmt_pct(as_float(ctx["latest"]["auto_payment_rate"]), 2), fmt_pp(as_float(ctx["latest"]["auto_payment_rate_change"]), 2)],
            ["처리효율", fmt_pct(as_float(ctx["latest"]["processing_efficiency"]), 2), fmt_pp(as_float(ctx["latest"]["processing_efficiency_change"]), 2)],
            ["운영상태", str(ctx["latest"].get("status_label", "")), "유지" if str(ctx["latest"].get("status_label", "")) == str(ctx["previous"].get("status_label", "")) else str(ctx["previous"].get("status_label", ""))],
        ],
        header_size=12,
        row_size=12,
    )
    add_blank_paragraph(doc)

    add_text_paragraph(
        doc,
        [("☐", {"font_name": FONT_SYMBOL}), (" ", {}), ("접수채널별 처리현황", {})],
        size=12,
    )
    add_note_paragraph(doc, ctx["channel_note"])
    add_table(
        doc,
        ["접수채널", "전체 청구건수", "자동지급건수", "자동지급 비율", "전월比"],
        [
            [
                row["channel"],
                fmt_int(row["total"]),
                fmt_int(row["auto"]),
                fmt_pct(row["share"], 1) if row["share"] is not None else "미계산",
                fmt_pp(row["delta_share"], 1),
            ]
            for row in ctx["channel_rows"]
        ],
        header_size=12,
        row_size=12,
    )
    add_blank_paragraph(doc)

    add_text_paragraph(
        doc,
        [("☐", {"font_name": FONT_SYMBOL}), (" ", {}), ("진료구분별 처리현황", {})],
        size=12,
    )
    add_note_paragraph(doc, ctx["treatment_note"])
    add_table(
        doc,
        ["진료구분", "자동지급건수", "비율", "전월比"],
        [
            [
                row["treatment"],
                fmt_int(row["count"]),
                fmt_pct(row["share"], 1) if row["share"] is not None else "미계산",
                fmt_pp(row["delta_share"], 1),
            ]
            for row in ctx["treatment_rows"]
        ],
        header_size=12,
        row_size=12,
    )

    add_text_paragraph(doc, [("2. ", {"bold": True}), ("유형별 상세 현황", {"bold": True})], size=12, bold=True)
    add_text_paragraph(
        doc,
        [("☐", {"font_name": FONT_SYMBOL}), (" ", {}), ("자동지급", {}), (" ", {}), ("유형", {})],
        size=12,
    )
    add_note_paragraph(doc, ctx["type_note"])
    add_type_table(doc, ctx["auto_type_rows"])
    add_blank_paragraph(doc)

    add_text_paragraph(
        doc,
        [("☐", {"font_name": FONT_SYMBOL}), (" ", {}), ("자동지급", {}), (" ", {}), ("제외", {}), (" ", {}), ("유형", {})],
        size=12,
    )
    add_text_paragraph(
        doc,
        [("- 자동지급 불가한 조건에 해당하여 인심사 배분되는 유형 중 전월대비 삭감/부지급 비율이 5%이상 낮아진 건", {})],
        size=12,
        line_spacing=1.5,
        first_line_indent=Inches(0.167),
    )
    add_note_paragraph(doc, ctx["exclusion_note"])
    add_type_table(doc, ctx["exclusion_rows"])
    add_blank_paragraph(doc)

    add_text_paragraph(doc, [("3. ", {"bold": True}), ("Action Item", {"bold": True})], size=12, bold=True)
    for line in ctx["next_actions"]:
        add_note_paragraph(doc, line)

    return doc


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", choices=["auto", "db", "local"], default="auto")
    parser.add_argument("--month", default=DEFAULT_REPORT_MONTH)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    snapshots, current_claims, previous_claims, auto_types, exclusion_types = load_data(args.source, args.month)
    ctx = build_context(snapshots, current_claims, previous_claims, auto_types, exclusion_types, args.month)

    doc = build_document(ctx)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)
    print(args.out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
